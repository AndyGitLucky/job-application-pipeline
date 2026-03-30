"""
generate_application.py
=======================
Generates application assets for selected jobs from jobs_scored.json.
"""

from __future__ import annotations

import argparse
import json
import logging
import re
import shutil
import time
from datetime import datetime
from pathlib import Path
from urllib.parse import urlparse

if __package__ in {None, ""}:
    import sys

    sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from source.candidate_profile import PROFILE_TEXT
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx2pdf import convert as docx2pdf_convert
from source.llm_client import llm_complete
from source.pipeline_state_manager import (
    attach_job_artifact,
    load_pipeline_state,
    save_pipeline_state,
    sync_jobs,
    update_job_stage,
)
from source.project_paths import ROOT_DIR, artifacts_path, resolve_artifacts_path, resolve_runtime_path, runtime_path
from source.retrieval_context import format_retrieval_context
from source.text_guardrails import find_negative_self_disclosure

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s  %(levelname)-8s  %(message)s",
    datefmt="%H:%M:%S",
)
log = logging.getLogger(__name__)

CONFIG = {
    "input_file": str(runtime_path("jobs_scored.json")),
    "output_dir": str(artifacts_path("applications")),
    "cover_letter_dir": str(ROOT_DIR / "Cover Letters"),
    "request_delay": 1.0,
    "min_score": 6,
    "allowed_decisions": {"apply"},
    "allowed_buckets": {"manual_apply_ready", "autoapply_ready"},
    "language": "de",
    "max_retries": 3,
    "guardrail_block_generation": True,
}

CANDIDATE = {
    "name": "Andreas Eichmann",
    "email": "andreas.eichmann@hotmail.com",
    "phone": "0176 3866 3585",
    "location": "München",
    "linkedin": "linkedin.com/in/andreas-eichmann",
    "github": "github.com/AndyGitLucky",
    "portfolio": "andygitlucky.github.io",
}

ANSCHREIBEN_PROMPT = """Du schreibst ein Anschreiben fuer eine Bewerbung.

KANDIDATENPROFIL:
{profile}

RELEVANTE KONTEXT-HINWEISE:
{retrieval_context}

STELLE:
Titel: {title}
Unternehmen: {company}
Beschreibung:
{description}

ANFORDERUNGEN:
- Sprache: {language}
- Ton: direkt, sachlich, selbstsicher
- Laenge: max. 250 Woerter
- Kein "Sehr geehrte Damen und Herren"
- Erster Absatz: Bezug zur Rolle/Firma
- Zweiter Absatz: konkrete Passung
- Dritter Absatz: 3 Stichpunkte
- Schluss: kurz und klar
- Nur Fliesstext, keine Adresszeilen
- Nenne NICHT ungefragt negative Selbstauskunft wie fehlenden Hochschulabschluss, Defizite oder Ausschlusskriterien
- Solche Risiken gehoeren NICHT ins Anschreiben, ausser die Stelle verlangt ausdruecklich eine direkte Stellungnahme dazu

Antworte NUR mit dem Anschreiben-Text.
"""

OUTREACH_PROMPT = """Du schreibst eine kurze LinkedIn/Email Outreach-Nachricht.

KANDIDATENPROFIL:
{profile}

RELEVANTE KONTEXT-HINWEISE:
{retrieval_context}

ZIELUNTERNEHMEN: {company}
ZIELROLLE DES EMPFAENGERS: {contact_role}
STELLE: {title}

ANFORDERUNGEN:
- Sprache: {language}
- Max. 60 Woerter
- Natuerlicher, menschlicher Ton
- Eine kurze Frage am Ende
- Keine generischen Floskeln
- Keine negativen Selbstrelativierungen oder Hinweise auf formale Defizite

Antworte NUR mit dem Nachrichtentext.
"""


def call_llm(prompt: str, quality: bool = False) -> str:
    return llm_complete(prompt, quality=quality)


def clean_generated_text(text: str) -> str:
    cleaned = text.strip()
    cleaned = re.sub(r"```(?:text|markdown)?", "", cleaned, flags=re.IGNORECASE)
    cleaned = cleaned.replace("```", "")
    cleaned = re.sub(r"\*\*(.*?)\*\*", r"\1", cleaned)
    cleaned = re.sub(r"__(.*?)__", r"\1", cleaned)
    cleaned = re.sub(r"^[ \t]*[-*][ \t]*", "- ", cleaned, flags=re.MULTILINE)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def assert_text_guardrails(text: str, asset_type: str) -> None:
    findings = find_negative_self_disclosure(text)
    if findings and CONFIG["guardrail_block_generation"]:
        raise ValueError(
            f"{asset_type} enthaelt unzulaessige Negativ-Selbstauskunft: {', '.join(sorted(set(findings)))}"
        )


def generate_anschreiben(job: dict) -> str:
    language = "Deutsch" if CONFIG["language"] == "de" else "Englisch"
    prompt = ANSCHREIBEN_PROMPT.format(
        profile=PROFILE_TEXT,
        retrieval_context=format_retrieval_context(
            job,
            mode="application",
            exclude_categories={"constraint"},
        ),
        title=job["title"],
        company=job["company"],
        description=job["description"][:1500],
        language=language,
    )
    return clean_generated_text(call_llm(prompt, quality=True))


def generate_outreach(job: dict, contact_role: str = "Hiring Manager") -> str:
    language = "Deutsch" if CONFIG["language"] == "de" else "Englisch"
    prompt = OUTREACH_PROMPT.format(
        profile=PROFILE_TEXT,
        retrieval_context=format_retrieval_context(
            job,
            mode="application",
            exclude_categories={"constraint"},
        ),
        company=job["company"],
        contact_role=contact_role,
        title=job["title"],
        language=language,
    )
    return clean_generated_text(call_llm(prompt))


def save_as_docx(anschreiben: str, job: dict, output_path: Path) -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    sender = doc.add_paragraph()
    sender.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = sender.add_run(
        f"{CANDIDATE['name']}\n"
        f"{CANDIDATE['location']}\n"
        f"{CANDIDATE['phone']}\n"
        f"{CANDIDATE['email']}\n"
        f"{CANDIDATE['linkedin']}\n"
        f"{CANDIDATE['github']}\n"
        f"{CANDIDATE['portfolio']}"
    )
    run.font.size = Pt(10)

    doc.add_paragraph()
    date_paragraph = doc.add_paragraph(datetime.today().strftime("%d.%m.%Y"))
    date_paragraph.runs[0].font.size = Pt(10)
    doc.add_paragraph()

    subject = doc.add_paragraph(f"Bewerbung als {job['title']}")
    subject.runs[0].bold = True
    subject.runs[0].font.size = Pt(11)
    doc.add_paragraph()

    for para in anschreiben.split("\n\n"):
        if not para.strip():
            continue
        paragraph = doc.add_paragraph(para.strip())
        paragraph.runs[0].font.size = Pt(11)
        paragraph.paragraph_format.space_after = Pt(8)

    doc.save(str(output_path))


def save_docx_as_pdf(docx_path: Path, pdf_path: Path) -> None:
    docx2pdf_convert(str(docx_path.resolve()), str(pdf_path.resolve()))


def keep_local_cover_letter_copy(central_pdf_path: Path, output_dir: Path) -> Path:
    local_pdf_path = output_dir / "anschreiben.pdf"
    if central_pdf_path.resolve() != local_pdf_path.resolve():
        shutil.copy2(central_pdf_path, local_pdf_path)
    return local_pdf_path


def make_output_dir(job: dict) -> Path:
    company_raw = _safe_company_name(job.get("company") or "")
    company_raw = re.sub(r"^\s*Arbeitgeber\s*:?\s*", "", company_raw, flags=re.IGNORECASE).strip()
    if not company_raw:
        host = urlparse(job.get("url") or "").netloc.replace("www.", "")
        company_raw = host or "unbekannt"

    safe_company = re.sub(r"[^\w\-]+", "_", company_raw)
    safe_company = re.sub(r"_+", "_", safe_company).strip("_-")[:60] or "unbekannt"
    output_dir = resolve_artifacts_path(CONFIG["output_dir"]) / f"{safe_company}_{job['id']}"
    output_dir.mkdir(parents=True, exist_ok=True)
    return output_dir


def _short_company_name(company: str) -> str:
    text = _safe_company_name(company)
    text = re.sub(r"^\s*Arbeitgeber\s*:?\s*", "", text, flags=re.IGNORECASE).strip()
    text = re.sub(
        r"\b(gmbh|mbh|ag|se|kg|gbr|inc|corp|corporation|ltd|llc|group|holding|holdings|deutschland)\b",
        "",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(r"[^\w\s-]+", " ", text, flags=re.UNICODE)
    parts = [part for part in re.split(r"[\s_-]+", text) if part]
    short = "_".join(parts[:3]) if parts else "Company"
    short = re.sub(r"_+", "_", short).strip("_")
    return short[:28] or "Company"


def _safe_company_name(company: str) -> str:
    text = str(company or "").strip()
    if text.lower() in {"nan", "none", "null", "unknown", "unbekannt"}:
        return ""
    return text


def make_cover_letter_pdf_path(job: dict) -> Path:
    cover_dir = Path(CONFIG["cover_letter_dir"])
    cover_dir.mkdir(parents=True, exist_ok=True)
    short_company = _short_company_name(job.get("company") or "")
    base_name = f"Andreas_Eichmann_CL_{short_company}"
    pdf_path = cover_dir / f"{base_name}.pdf"
    if not pdf_path.exists():
        return pdf_path

    counter = 2
    while True:
        candidate = cover_dir / f"{base_name}_{counter}.pdf"
        if not candidate.exists():
            return candidate
        counter += 1


def generate_applications(
    input_file: str | None = None,
    *,
    force: bool = False,
    job_ids: set[str] | None = None,
    limit: int | None = None,
) -> list:
    input_path = resolve_runtime_path(input_file or CONFIG["input_file"])
    if not input_path.exists():
        log.error("Input-Datei nicht gefunden: %s", input_path)
        return []

    jobs = json.loads(input_path.read_text(encoding="utf-8"))
    state = load_pipeline_state()
    sync_jobs(state, jobs, stage="generation")

    retry_blocked = {
        job_id for job_id, job_state in state.get("jobs", {}).items()
        if int(job_state.get("retry_count", 0)) >= CONFIG["max_retries"]
    }

    allowed_decisions = set(CONFIG["allowed_decisions"])

    if job_ids:
        to_process = [
            job for job in jobs
            if str(job.get("id", "")).strip() in job_ids
            and str(job.get("id", "")) not in retry_blocked
            and job.get("job_status", "live") != "invalid"
            and (force or not job.get("application_generated"))
            and (force or job.get("decision") in allowed_decisions)
            and (force or job.get("final_bucket") in CONFIG["allowed_buckets"])
        ]
    else:
        to_process = [
            job for job in jobs
            if job.get("recommended")
            and job.get("score", 0) >= CONFIG["min_score"]
            and job.get("decision") in allowed_decisions
            and job.get("final_bucket") in CONFIG["allowed_buckets"]
            and (force or not job.get("application_generated"))
            and str(job.get("id", "")) not in retry_blocked
            and job.get("job_status", "live") != "invalid"
        ]

    if limit is not None:
        to_process = to_process[: max(0, int(limit))]

    log.info("Jobs geladen: %s gesamt, %s zu verarbeiten", len(jobs), len(to_process))
    generated = []

    for idx, job in enumerate(to_process, start=1):
        log.info("[%s/%s] %s @ %s", idx, len(to_process), job["title"], job["company"])
        update_job_stage(state, job["id"], "generation", "in_progress", message="application_generation_started")
        try:
            output_dir = make_output_dir(job)
            anschreiben = generate_anschreiben(job)
            assert_text_guardrails(anschreiben, "anschreiben")
            time.sleep(CONFIG["request_delay"])
            outreach = generate_outreach(job)
            assert_text_guardrails(outreach, "outreach")
            time.sleep(CONFIG["request_delay"])

            anschreiben_txt = output_dir / "anschreiben.txt"
            outreach_txt = output_dir / "outreach.txt"
            meta_path = output_dir / "meta.json"
            docx_path = output_dir / "anschreiben.docx"
            pdf_path = make_cover_letter_pdf_path(job)

            anschreiben_txt.write_text(anschreiben, encoding="utf-8")
            outreach_txt.write_text(outreach, encoding="utf-8")
            save_as_docx(anschreiben, job, docx_path)
            save_docx_as_pdf(docx_path, pdf_path)
            local_pdf_path = keep_local_cover_letter_copy(pdf_path, output_dir)
        except Exception as exc:
            log.warning("Generation failed for %s: %s", job["id"], exc)
            update_job_stage(state, job["id"], "generation", "failed", error=str(exc), message="generation_failed")
            continue

        meta = {
            "job_id": job["id"],
            "title": job["title"],
            "company": job["company"],
            "url": job["url"],
            "score": job["score"],
            "decision": job.get("decision"),
            "contact_email": job.get("contact_email", ""),
            "contact_name": job.get("contact_name", ""),
            "contact_role": job.get("contact_role", ""),
            "cover_letter_pdf": str(pdf_path),
            "cover_letter_pdf_local": str(local_pdf_path),
            "generated": datetime.now().isoformat(),
        }
        meta_path.write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")

        job["application_generated"] = True
        job["application_dir"] = str(output_dir)
        job["cover_letter_pdf"] = str(pdf_path)
        job["cover_letter_pdf_local"] = str(local_pdf_path)
        generated.append(job)

        attach_job_artifact(state, job["id"], "application_dir", str(output_dir))
        attach_job_artifact(state, job["id"], "cover_letter_txt", str(anschreiben_txt))
        attach_job_artifact(state, job["id"], "outreach_txt", str(outreach_txt))
        if docx_path.exists():
            attach_job_artifact(state, job["id"], "cover_letter_docx", str(docx_path))
        if pdf_path.exists():
            attach_job_artifact(state, job["id"], "cover_letter_pdf", str(pdf_path))
        if local_pdf_path.exists():
            attach_job_artifact(state, job["id"], "cover_letter_pdf_local", str(local_pdf_path))
        update_job_stage(
            state,
            job["id"],
            "generation",
            "completed",
            message="application_assets_created",
            extras={"application_dir": str(output_dir)},
        )

    input_path.write_text(json.dumps(jobs, ensure_ascii=False, indent=2), encoding="utf-8")
    save_pipeline_state(state)
    log.info("OK. %s Bewerbungen generiert unter ./%s/", len(generated), CONFIG["output_dir"])
    return generated


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generiert Bewerbungsunterlagen aus jobs_scored.json")
    parser.add_argument("--input", default=None)
    parser.add_argument("--force", action="store_true")
    parser.add_argument("--job-ids", default="")
    parser.add_argument("--limit", type=int, default=None)
    args = parser.parse_args()

    job_id_set = {item.strip() for item in (args.job_ids or "").split(",") if item.strip()} or None
    generated_jobs = generate_applications(args.input, force=args.force, job_ids=job_id_set, limit=args.limit)
    print(f"\nOK. {len(generated_jobs)} Bewerbungen generiert.")
